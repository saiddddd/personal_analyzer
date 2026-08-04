import os
import re
import sys
import uuid

try:
    __import__("pysqlite3")
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    pass

import streamlit as st
import chromadb
import pandas as pd
from chromadb.utils import embedding_functions
from groq import Groq
from pypdf import PdfReader

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import requests
    from bs4 import BeautifulSoup
    SCRAPE_AVAILABLE = True
except ImportError:
    SCRAPE_AVAILABLE = False

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

st.set_page_config(page_title="RAG Dashboard", page_icon="📚", layout="wide")


def get_secret_api_key():
    """Cek apakah API key sudah dikonfigurasi lewat Streamlit Secrets."""
    try:
        key = st.secrets["GROQ_API_KEY"]
        return key if key else None
    except Exception:
        return None


# ============================================================
# SIDEBAR - KONFIGURASI
# ============================================================
st.sidebar.title("⚙️ Konfigurasi")

_secret_key = get_secret_api_key()
if _secret_key:
    # Sudah dikonfigurasi admin lewat Secrets -> gak perlu tampilkan kolom input sama sekali
    groq_api_key = _secret_key
    st.sidebar.success("✅ Groq API Key sudah dikonfigurasi.")
else:
    groq_api_key = st.sidebar.text_input(
        "Groq API Key",
        type="password",
        value=os.getenv("GROQ_API_KEY", ""),
        help="Ambil gratis di https://console.groq.com/keys",
    )
    st.sidebar.markdown("[🔑 Ambil API Key gratis di sini](https://console.groq.com/keys)")

model_name = st.sidebar.selectbox(
    "Model LLM (Groq)",
    ["llama-3.3-70b-versatile", "llama-3.1-8b-instant", "gemma2-9b-it"],
    help="70b lebih pintar tapi sedikit lebih lambat, 8b paling cepat.",
)

chunk_size = st.sidebar.slider("Chunk size (karakter)", 200, 2000, 800, step=100)
chunk_overlap = st.sidebar.slider("Chunk overlap (karakter)", 0, 500, 100, step=50)
top_k = st.sidebar.slider("Jumlah chunk yang diambil (top-k)", 1, 10, 3)
history_turns = st.sidebar.slider("Riwayat chat diikutkan (turn)", 0, 5, 2)

st.sidebar.markdown("---")
st.sidebar.subheader("🤖 Agentic RAG")
agentic_mode = st.sidebar.checkbox(
    "Aktifkan query decomposition + auto re-retrieval",
    value=True,
    help="Pertanyaan kompleks otomatis dipecah jadi sub-pertanyaan, dan pencarian yang skornya rendah otomatis dicoba ulang dengan query alternatif.",
)
relevance_threshold = st.sidebar.slider(
    "Ambang batas re-retrieval (skor relevansi)", 0.0, 1.0, 0.3, step=0.05,
)

if st.sidebar.button("🗑️ Reset Semua Data"):
    for key in ["messages", "docs_loaded", "doc_meta", "collection_name", "dataframes", "data_messages"]:
        st.session_state.pop(key, None)
    st.cache_resource.clear()
    st.rerun()

with st.sidebar.expander("ℹ️ Status fitur opsional"):
    st.write(f"{'✅' if DOCX_AVAILABLE else '❌'} DOCX support")
    st.write(f"{'✅' if SCRAPE_AVAILABLE else '❌'} Website scraping")
    st.write(f"{'✅' if OCR_AVAILABLE else '❌'} OCR untuk PDF scan")
    if not OCR_AVAILABLE:
        st.caption("Install `pytesseract` + `pdf2image` (pip) dan Tesseract-OCR + Poppler (sistem) untuk aktifkan OCR.")

# ============================================================
# INISIALISASI VECTOR DB (ChromaDB, PERSISTEN di disk)
# ============================================================
@st.cache_resource
def get_chroma_client():
    return chromadb.PersistentClient(path="./chroma_db")


@st.cache_resource
def get_embedding_fn():
    return embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name="all-MiniLM-L6-v2"
    )


client = get_chroma_client()
embed_fn = get_embedding_fn()

if "collection_name" not in st.session_state:
    st.session_state.collection_name = "rag_dashboard_kb"

collection = client.get_or_create_collection(
    name=st.session_state.collection_name,
    embedding_function=embed_fn,
    metadata={"hnsw:space": "cosine"},
)

if "messages" not in st.session_state:
    st.session_state.messages = []
if "docs_loaded" not in st.session_state:
    st.session_state.docs_loaded = []
if "dataframes" not in st.session_state:
    st.session_state.dataframes = {}
if "data_messages" not in st.session_state:
    st.session_state.data_messages = []
if "doc_meta" not in st.session_state:
    st.session_state.doc_meta = {}
    try:
        existing = collection.get(include=["metadatas"])
        for m in existing["metadatas"]:
            src = m.get("source")
            if src and src not in st.session_state.doc_meta:
                st.session_state.doc_meta[src] = {"pages": m.get("total_pages"), "chunks": 0}
            if src:
                st.session_state.doc_meta[src]["chunks"] += 1
        st.session_state.docs_loaded = [
            f"{name} — {meta['chunks']} chunks" + (f", {meta['pages']} halaman" if meta.get("pages") else "")
            for name, meta in st.session_state.doc_meta.items()
        ]
    except Exception:
        pass


# ============================================================
# HELPER FUNCTIONS - DOKUMEN (RAG)
# ============================================================
def extract_text(file):
    """Ambil teks dari PDF (+ OCR fallback), DOCX, atau TXT. Return (teks, jumlah_halaman)."""
    name = file.name.lower()

    if file.type == "application/pdf" or name.endswith(".pdf"):
        file.seek(0)
        reader = PdfReader(file)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        num_pages = len(reader.pages)

        if len(text.strip()) < 50 and OCR_AVAILABLE:
            st.info(f"'{file.name}' terdeteksi minim teks, mencoba OCR...")
            try:
                file.seek(0)
                images = convert_from_bytes(file.read())
                ocr_parts = []
                for img in images:
                    try:
                        ocr_parts.append(pytesseract.image_to_string(img, lang="ind+eng"))
                    except Exception:
                        ocr_parts.append(pytesseract.image_to_string(img, lang="eng"))
                ocr_text = "\n".join(ocr_parts)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
            except Exception as e:
                st.warning(f"OCR gagal untuk '{file.name}': {e}")
        return text, num_pages

    if name.endswith(".docx"):
        if not DOCX_AVAILABLE:
            st.error("Package `python-docx` belum terinstall. Jalankan: pip install python-docx")
            return "", None
        file.seek(0)
        document = docx.Document(file)
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts), None

    file.seek(0)
    return file.read().decode("utf-8", errors="ignore"), None


def chunk_text(text: str, size: int, overlap: int) -> list[str]:
    """Chunking yang menghormati batas paragraf (tidak motong kata di tengah)."""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n|\n", text) if p.strip()]
    raw_chunks, current = [], ""
    for para in paragraphs:
        if len(current) + len(para) + 1 <= size:
            current = (current + "\n" + para).strip()
        else:
            if current:
                raw_chunks.append(current)
            if len(para) > size:
                step = max(size - overlap, 1)
                start = 0
                while start < len(para):
                    raw_chunks.append(para[start:start + size].strip())
                    start += step
                current = ""
            else:
                current = para
    if current:
        raw_chunks.append(current)

    if overlap > 0 and len(raw_chunks) > 1:
        overlapped = [raw_chunks[0]]
        for i in range(1, len(raw_chunks)):
            tail = raw_chunks[i - 1][-overlap:]
            overlapped.append((tail + " " + raw_chunks[i]).strip())
        raw_chunks = overlapped

    return [c for c in raw_chunks if c.strip()]


def add_document(file):
    text, num_pages = extract_text(file)

    if len(text.strip()) < 50:
        hint = "" if OCR_AVAILABLE else " (OCR belum aktif — install pytesseract+pdf2image kalau ini PDF hasil scan)"
        st.warning(f"⚠️ '{file.name}' teksnya sangat sedikit/kosong.{hint}")
        return

    chunks = chunk_text(text, chunk_size, chunk_overlap)
    if not chunks:
        return

    ids = [f"{file.name}_{i}_{uuid.uuid4().hex[:6]}" for i in range(len(chunks))]
    metadatas = [
        {"source": file.name, "chunk": i, "total_pages": num_pages or 0}
        for i in range(len(chunks))
    ]
    collection.add(documents=chunks, ids=ids, metadatas=metadatas)

    st.session_state.doc_meta[file.name] = {"pages": num_pages, "chunks": len(chunks)}
    info = f"{file.name} — {len(chunks)} chunks"
    if num_pages:
        info += f", {num_pages} halaman"
    st.session_state.docs_loaded.append(info)


def add_website(url: str):
    if not SCRAPE_AVAILABLE:
        st.error("Package `requests` dan `beautifulsoup4` belum terinstall.")
        return
    try:
        resp = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
    except Exception as e:
        st.error(f"Gagal mengambil {url}: {e}")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "svg"]):
        tag.decompose()
    text = re.sub(r"\n\s*\n+", "\n\n", soup.get_text(separator="\n")).strip()

    if len(text) < 50:
        st.warning(f"Konten dari {url} minim/kosong (mungkin butuh JavaScript rendering yang tidak didukung scraper sederhana ini).")
        return

    chunks = chunk_text(text, chunk_size, chunk_overlap)
    ids = [f"{url}_{i}_{uuid.uuid4().hex[:6]}" for i in range(len(chunks))]
    metadatas = [{"source": url, "chunk": i, "total_pages": 0} for i in range(len(chunks))]
    collection.add(documents=chunks, ids=ids, metadatas=metadatas)
    st.session_state.doc_meta[url] = {"pages": None, "chunks": len(chunks)}
    st.session_state.docs_loaded.append(f"{url} — {len(chunks)} chunks (website)")


def detect_meta_question(query: str):
    q = query.lower()
    if any(k in q for k in ["berapa halaman", "jumlah halaman", "total halaman", "how many pages"]):
        return "pages"
    if any(k in q for k in ["dokumen apa saja", "list dokumen", "nama file", "nama dokumen", "ada dokumen apa", "berapa dokumen"]):
        return "list"
    return None


def answer_meta_question(kind: str) -> str:
    if not st.session_state.doc_meta:
        return "Belum ada dokumen yang diupload."
    if kind == "pages":
        lines = []
        for name, meta in st.session_state.doc_meta.items():
            pages = meta.get("pages")
            lines.append(f"- **{name}**: {pages if pages else 'tidak diketahui'} halaman")
        return "Jumlah halaman dokumen:\n" + "\n".join(lines)
    if kind == "list":
        lines = [f"- **{name}** ({meta['chunks']} chunks)" for name, meta in st.session_state.doc_meta.items()]
        return "Dokumen yang ada di knowledge base:\n" + "\n".join(lines)
    return "Tidak bisa menjawab pertanyaan ini."


def retrieve_context(query: str, k: int):
    results = collection.query(query_texts=[query], n_results=k)
    docs = results["documents"][0] if results["documents"] else []
    metas = results["metadatas"][0] if results["metadatas"] else []
    dists = results["distances"][0] if results.get("distances") else []
    ids = results["ids"][0] if results.get("ids") else []
    return docs, metas, dists, ids


def decompose_query(query: str, api_key: str, model: str) -> list[str]:
    groq_client = Groq(api_key=api_key)
    system_prompt = (
        "Kamu membantu memecah pertanyaan pengguna menjadi sub-pertanyaan yang lebih fokus "
        "untuk pencarian dokumen (retrieval). Jika pertanyaan sudah simpel dan satu topik, "
        "kembalikan pertanyaan itu apa adanya, satu baris saja. Jika pertanyaan kompleks, "
        "pecah jadi maksimal 3 sub-pertanyaan singkat, masing-masing di baris terpisah diawali '- '. "
        "JANGAN menjawab pertanyaannya, JANGAN tambahkan penjelasan apa pun."
    )
    try:
        resp = groq_client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": query}],
            temperature=0,
            max_tokens=200,
        )
        text = resp.choices[0].message.content or ""
        lines = [l.strip().lstrip("-").strip() for l in text.splitlines() if l.strip()]
        lines = [l for l in lines if l]
        return lines[:3] if lines else [query]
    except Exception:
        return [query]


def rephrase_query(query: str, api_key: str, model: str) -> str:
    groq_client = Groq(api_key=api_key)
    system_prompt = (
        "Tulis ulang pertanyaan berikut jadi satu query pencarian alternatif dengan kata kunci "
        "berbeda, tapi maksud tetap sama. Jawab HANYA query barunya, tanpa penjelasan."
    )
    try:
        resp = groq_client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": query}],
            temperature=0.3,
            max_tokens=60,
        )
        alt = (resp.choices[0].message.content or "").strip().strip('"')
        return alt if alt and alt.lower() != query.lower() else ""
    except Exception:
        return ""


def agentic_retrieve(query: str, k: int, api_key: str, model: str, use_agentic: bool, threshold: float):
    log = []
    subqueries = decompose_query(query, api_key, model) if use_agentic else [query]

    if len(subqueries) > 1:
        log.append(f"🔍 Pertanyaan dipecah jadi {len(subqueries)} sub-pertanyaan: " + " | ".join(subqueries))
    else:
        log.append("🔍 Pertanyaan simpel, retrieval langsung tanpa decomposition.")

    collected = {}
    for sq in subqueries:
        docs, metas, dists, ids = retrieve_context(sq, k)
        best_score = max([1 - d for d in dists], default=0.0)

        if use_agentic and best_score < threshold:
            log.append(f"⚠️ Skor relevansi rendah ({best_score:.2f}) untuk '{sq}' → mencoba re-retrieval.")
            alt = rephrase_query(sq, api_key, model)
            if alt:
                log.append(f"🔁 Query alternatif: '{alt}'")
                docs2, metas2, dists2, ids2 = retrieve_context(alt, k)
                docs, metas, dists, ids = docs + docs2, metas + metas2, dists + dists2, ids + ids2

        for doc, meta, dist, id_ in zip(docs, metas, dists, ids):
            if id_ not in collected or dist < collected[id_][2]:
                collected[id_] = (doc, meta, dist)

    sorted_items = sorted(collected.values(), key=lambda x: x[2])[: max(k, 5)]
    final_docs = [it[0] for it in sorted_items]
    final_metas = [it[1] for it in sorted_items]
    final_dists = [it[2] for it in sorted_items]
    log.append(f"✅ Total {len(final_docs)} chunk unik dipakai sebagai konteks akhir.")
    return final_docs, final_metas, final_dists, log


def build_history_snippet(n_turns: int) -> str:
    if n_turns <= 0:
        return ""
    recent = st.session_state.messages[-(n_turns * 2):]
    lines = [f"{'User' if m['role']=='user' else 'Asisten'}: {m['content']}" for m in recent]
    return "\n".join(lines)


def ask_groq_stream(query: str, context_docs: list[str], history: str, api_key: str, model: str):
    groq_client = Groq(api_key=api_key)
    context = "\n\n---\n\n".join(context_docs) if context_docs else "(tidak ada konteks relevan ditemukan)"
    system_prompt = (
        "Kamu adalah asisten yang menjawab HANYA berdasarkan konteks dokumen yang diberikan. "
        "Jika jawabannya tidak ada di dalam konteks, katakan dengan jujur bahwa kamu tidak tahu, "
        "jangan mengarang. Selalu jawab dalam Bahasa Indonesia yang jelas dan ringkas."
    )
    user_prompt = ""
    if history:
        user_prompt += f"Riwayat percakapan sebelumnya:\n{history}\n\n"
    user_prompt += f"Konteks dokumen:\n{context}\n\nPertanyaan: {query}"

    try:
        stream = groq_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
            stream=True,
        )
        for chunk in stream:
            delta = chunk.choices[0].delta.content
            if delta:
                yield delta
    except Exception as e:
        msg = str(e)
        if "rate limit" in msg.lower() or "429" in msg:
            yield "⚠️ Kena rate limit Groq (kuota gratis habis sementara). Coba tunggu sebentar lalu ulangi."
        elif "401" in msg or "auth" in msg.lower():
            yield "⚠️ API Key Groq tidak valid. Cek lagi di sidebar."
        else:
            yield f"⚠️ Terjadi error saat memanggil Groq API: {msg}"


# ============================================================
# HELPER FUNCTIONS - ANALISIS DATA (CSV/EXCEL)
# ============================================================
def is_code_safe(code: str) -> bool:
    banned = [
        "import ", "open(", "exec(", "eval(", "os.", "sys.", "subprocess",
        "__", "globals(", "locals(", "input(", "compile(", "getattr(",
        "setattr(", "delattr(", "system(", "popen(",
    ]
    low = code.lower()
    return not any(b in low for b in banned)


def generate_pandas_code(question: str, df_name: str, df: "pd.DataFrame", api_key: str, model: str, error_hint: str = None) -> str:
    groq_client = Groq(api_key=api_key)
    schema = (
        f"Nama file: {df_name}\n"
        f"Kolom & tipe data:\n{df.dtypes.to_string()}\n\n"
        f"5 baris contoh:\n{df.head(5).to_string()}"
    )
    system_prompt = (
        "Kamu asisten data analyst yang menulis kode pandas Python. "
        f"Tersedia dataframe bernama `df` dengan struktur:\n{schema}\n\n"
        "Tulis kode python singkat yang menghitung jawaban dari pertanyaan user dan "
        "menyimpan hasil akhirnya ke variabel bernama `result`. Hanya boleh pakai `df` "
        "dan `pd` (pandas). JANGAN import apa pun, JANGAN pakai open/exec/eval/os/sys/"
        "subprocess/input, JANGAN akses file atau internet. Jawab HANYA dengan kode "
        "python, tanpa markdown, tanpa backtick, tanpa penjelasan."
    )
    user_prompt = f"Pertanyaan: {question}"
    if error_hint:
        user_prompt += f"\n\nKode sebelumnya error dengan pesan: {error_hint}\nPerbaiki kodenya."

    resp = groq_client.chat.completions.create(
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0,
        max_tokens=300,
    )
    code = (resp.choices[0].message.content or "").strip()
    code = re.sub(r"^```(python)?|```$", "", code, flags=re.MULTILINE).strip()
    return code


def run_pandas_code(code: str, df: "pd.DataFrame"):
    safe_globals = {"__builtins__": {}, "pd": pd}
    safe_locals = {"df": df}
    exec(code, safe_globals, safe_locals)
    return safe_locals.get("result", safe_globals.get("result"))


def explain_result(question: str, result, api_key: str, model: str) -> str:
    groq_client = Groq(api_key=api_key)
    system_prompt = (
        "Jelaskan hasil perhitungan data berikut secara natural dan ringkas dalam Bahasa "
        "Indonesia. Beri sedikit insight, jangan cuma mengulang angkanya doang."
    )
    user_prompt = f"Pertanyaan: {question}\nHasil: {result}"
    try:
        resp = groq_client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            temperature=0.3,
            max_tokens=300,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"(Hasil: {result}) — gagal membuat penjelasan natural: {e}"


# ============================================================
# MAIN UI
# ============================================================
st.title("📚 RAG Dashboard (Groq + ChromaDB + Streamlit)")
st.caption("Tanya dokumen (PDF/DOCX/TXT/Website) atau analisis data (CSV/Excel). 100% pakai tools gratis.")

tab_doc, tab_data = st.tabs(["📄 Tanya Dokumen (RAG)", "📊 Analisis Data (CSV/Excel)"])

# ---------------- TAB 1: DOKUMEN (RAG) ----------------
with tab_doc:
    col1, col2 = st.columns([1, 2])

    with col1:
        st.subheader("1️⃣ Upload Dokumen")
        uploaded_files = st.file_uploader(
            "Upload PDF, DOCX, atau TXT (bisa lebih dari satu)",
            type=["pdf", "txt", "docx"],
            accept_multiple_files=True,
        )
        if uploaded_files and st.button("🚀 Proses Dokumen", use_container_width=True):
            with st.spinner("Memproses & meng-embed dokumen..."):
                for f in uploaded_files:
                    add_document(f)
            st.success("Dokumen berhasil diproses dan siap ditanya!")

        st.markdown("**🌐 Atau tambah dari website**")
        website_url = st.text_input("URL website", placeholder="https://contoh.com/artikel")
        if st.button("🌐 Scrape & Tambah", use_container_width=True) and website_url:
            with st.spinner(f"Mengambil konten dari {website_url}..."):
                add_website(website_url)

        if st.session_state.docs_loaded:
            st.write("**📂 Dokumen di knowledge base:**")
            for d in st.session_state.docs_loaded:
                st.write(f"- {d}")
        else:
            st.info("Belum ada dokumen. Upload atau scrape website dulu.")

    with col2:
        st.subheader("2️⃣ Tanya Jawab")

        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                if msg.get("sources"):
                    with st.expander("📄 Sumber jawaban"):
                        for s in msg["sources"]:
                            st.caption(s)

        query = st.chat_input("Tanya sesuatu tentang dokumen kamu...", key="doc_chat")
        if query:
            st.session_state.messages.append({"role": "user", "content": query})
            with st.chat_message("user"):
                st.markdown(query)

            meta_kind = detect_meta_question(query)

            with st.chat_message("assistant"):
                if meta_kind:
                    answer = answer_meta_question(meta_kind)
                    st.markdown(answer)
                    sources = []
                elif not groq_api_key:
                    answer = "Isi Groq API Key dulu di sidebar ya."
                    st.error(answer)
                    sources = []
                elif collection.count() == 0:
                    answer = "Upload dokumen atau scrape website dulu sebelum bertanya."
                    st.error(answer)
                    sources = []
                else:
                    with st.spinner("Mencari konteks..."):
                        docs, metas, dists, agent_log = agentic_retrieve(
                            query, top_k, groq_api_key, model_name, agentic_mode, relevance_threshold
                        )
                        history = build_history_snippet(history_turns)
                        sources = [
                            f"{m['source']} (chunk {m['chunk']}, skor relevansi={1 - d:.2f})"
                            for m, d in zip(metas, dists)
                        ]

                    answer = st.write_stream(ask_groq_stream(query, docs, history, groq_api_key, model_name))

                    if agentic_mode:
                        with st.expander("🧠 Proses Agentic"):
                            for step in agent_log:
                                st.caption(step)

                    if sources:
                        with st.expander("📄 Sumber jawaban"):
                            for s, d in zip(sources, docs):
                                st.caption(s)
                                st.text(d[:300] + ("..." if len(d) > 300 else ""))

            st.session_state.messages.append({"role": "assistant", "content": answer, "sources": sources})

# ---------------- TAB 2: ANALISIS DATA (CSV/EXCEL) ----------------
with tab_data:
    col1, col2 = st.columns([1, 2])

    with col1:
        st.subheader("1️⃣ Upload Data")
        data_files = st.file_uploader(
            "Upload CSV atau Excel (bisa lebih dari satu)",
            type=["csv", "xlsx", "xls"],
            accept_multiple_files=True,
            key="data_uploader",
        )
        if data_files and st.button("📥 Muat Data", use_container_width=True):
            for f in data_files:
                try:
                    df_new = pd.read_csv(f) if f.name.lower().endswith(".csv") else pd.read_excel(f)
                    st.session_state.dataframes[f.name] = df_new
                    st.success(f"'{f.name}' dimuat: {df_new.shape[0]} baris x {df_new.shape[1]} kolom")
                except Exception as e:
                    st.error(f"Gagal memuat '{f.name}': {e}")

        if not st.session_state.dataframes:
            st.info("Belum ada data. Upload CSV/Excel dulu.")

    with col2:
        st.subheader("2️⃣ Analisis")
        if not st.session_state.dataframes:
            st.info("Upload data di panel kiri untuk mulai bertanya.")
        else:
            selected_df_name = st.selectbox("Pilih dataset", list(st.session_state.dataframes.keys()))
            df = st.session_state.dataframes[selected_df_name]

            with st.expander("👀 Preview data"):
                st.dataframe(df.head(20), use_container_width=True)
                st.caption(f"{df.shape[0]} baris, {df.shape[1]} kolom")

            for msg in st.session_state.data_messages:
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])

            data_query = st.chat_input(
                "Tanya soal data kamu (misal: rata-rata kolom X, top 5 by Y)...", key="data_chat"
            )
            if data_query:
                st.session_state.data_messages.append({"role": "user", "content": data_query})
                with st.chat_message("user"):
                    st.markdown(data_query)

                with st.chat_message("assistant"):
                    if not groq_api_key:
                        answer = "Isi Groq API Key dulu di sidebar ya."
                        st.error(answer)
                    else:
                        with st.spinner("Menyusun & menjalankan analisis..."):
                            code = generate_pandas_code(data_query, selected_df_name, df, groq_api_key, model_name)
                            result, error_msg = None, None

                            if not is_code_safe(code):
                                error_msg = "Kode yang dihasilkan mengandung operasi yang tidak diizinkan."
                            else:
                                try:
                                    result = run_pandas_code(code, df)
                                except Exception as e:
                                    error_msg = str(e)
                                    code_retry = generate_pandas_code(
                                        data_query, selected_df_name, df, groq_api_key, model_name, error_hint=error_msg
                                    )
                                    if is_code_safe(code_retry):
                                        try:
                                            result = run_pandas_code(code_retry, df)
                                            code, error_msg = code_retry, None
                                        except Exception as e2:
                                            error_msg = str(e2)

                        if error_msg:
                            answer = f"⚠️ Gagal menjalankan analisis: {error_msg}"
                            st.error(answer)
                            with st.expander("🔧 Kode yang dicoba"):
                                st.code(code, language="python")
                        else:
                            narrative = explain_result(data_query, result, groq_api_key, model_name)
                            st.markdown(narrative)
                            answer = narrative
                            with st.expander("🔧 Kode & hasil mentah"):
                                st.code(code, language="python")
                                if isinstance(result, (pd.Series, pd.DataFrame)):
                                    st.dataframe(result, use_container_width=True)
                                    if isinstance(result, pd.Series) and pd.api.types.is_numeric_dtype(result) and len(result) <= 30:
                                        st.bar_chart(result)
                                else:
                                    st.write(result)

                st.session_state.data_messages.append({"role": "assistant", "content": answer})
